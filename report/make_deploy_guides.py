#!/usr/bin/env python3
"""Generate two focused VAST-styled deployment guides:
   - KVCache-Demo-Automated-Deployment.docx  (deploy.sh path, ~8 min)
   - KVCache-Demo-Manual-Deployment.docx     (every component by hand)
"""
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

DEEP_BLUE = RGBColor(0x1E, 0x3A, 0x5F)
DEEP_BLUE_HEX = "1E3A5F"
LIGHT_BLUE_HEX = "F0F7FA"
CODE_BG = "F5F5F5"
BORDER_GRAY = "CCCCCC"
LOGO = ("/Users/wongtran/Library/Application Support/Claude/local-agent-mode-sessions/"
        "skills-plugin/3928ab7b-2926-439f-8a2e-ea3d7739d43a/cb212373-0fdf-4237-9256-3eefc9914e53/"
        "skills/vast-docx-style/assets/vast_logo.png")


class GuideDoc:
    def __init__(self, title, subtitle, meta_line):
        self.doc = Document()
        d = self.doc
        d.styles["Normal"].font.name = "Arial"
        d.styles["Normal"].font.size = Pt(11)
        for lvl, sz, before, after in (("Heading 1", 17, 14, 7), ("Heading 2", 13, 10, 5)):
            st = d.styles[lvl]
            st.font.name = "Arial"
            st.font.size = Pt(sz)
            st.font.bold = True
            st.font.color.rgb = DEEP_BLUE
            st.paragraph_format.space_before = Pt(before)
            st.paragraph_format.space_after = Pt(after)
        sec = d.sections[0]
        sec.page_width, sec.page_height = Inches(8.5), Inches(11)
        sec.left_margin = sec.right_margin = sec.top_margin = sec.bottom_margin = Inches(1)
        hdr = sec.header.paragraphs[0]
        hdr.paragraph_format.space_after = Pt(0)
        hdr.add_run().add_picture(LOGO, width=Inches(0.5))
        t = d.add_paragraph()
        t.paragraph_format.space_before = Pt(18)
        r = t.add_run(title)
        r.font.size = Pt(22)
        r.bold = True
        r.font.color.rgb = DEEP_BLUE
        s = d.add_paragraph()
        r = s.add_run(subtitle)
        r.font.size = Pt(12.5)
        r.font.color.rgb = RGBColor(0x55, 0x66, 0x77)
        m = d.add_paragraph()
        r = m.add_run(meta_line)
        r.font.size = Pt(9)
        r.italic = True
        r.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    def h1(self, text):
        self.doc.add_heading(text, level=1)

    def h2(self, text):
        self.doc.add_heading(text, level=2)

    def body(self, text, bold_lead=None, after=6, italic=False):
        p = self.doc.add_paragraph()
        p.paragraph_format.space_after = Pt(after)
        if bold_lead:
            r = p.add_run(bold_lead)
            r.bold = True
        r = p.add_run(text)
        r.italic = italic

    def bullet(self, text, bold_lead=None):
        p = self.doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(3)
        if bold_lead:
            r = p.add_run(bold_lead)
            r.bold = True
        p.add_run(text)

    def code(self, lines, after=8):
        if isinstance(lines, str):
            lines = lines.split("\n")
        for i, line in enumerate(lines):
            p = self.doc.add_paragraph()
            pPr = p._p.get_or_add_pPr()
            pPr.append(p._p.makeelement(qn("w:shd"), {qn("w:val"): "clear", qn("w:fill"): CODE_BG}))
            p.paragraph_format.space_after = Pt(after if i == len(lines) - 1 else 0)
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.left_indent = Inches(0.12)
            r = p.add_run(line if line else " ")
            r.font.name = "Consolas"
            r._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")
            r._element.rPr.rFonts.set(qn("w:cs"), "Consolas")
            r.font.size = Pt(9)

    def callout(self, label, text):
        p = self.doc.add_paragraph()
        pPr = p._p.get_or_add_pPr()
        pPr.append(p._p.makeelement(qn("w:shd"), {qn("w:val"): "clear", qn("w:fill"): LIGHT_BLUE_HEX}))
        p.paragraph_format.space_after = Pt(8)
        p.paragraph_format.left_indent = Inches(0.12)
        r = p.add_run(label + "  ")
        r.bold = True
        r.font.color.rgb = DEEP_BLUE
        p.add_run(text)

    def image(self, path, width=6.4, caption=None):
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(path, width=Inches(width))
        if caption:
            c = self.doc.add_paragraph()
            c.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = c.add_run(caption)
            r.italic = True
            r.font.size = Pt(9)
            r.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    def save(self, path):
        self.doc.save(path)
        print("saved", path)


# ==================================================================
# GUIDE 1 — AUTOMATED DEPLOYMENT
# ==================================================================
g = GuideDoc(
    "KV Cache Tiering Demo — Automated Deployment",
    "One command from an empty vast.ai account to a live interactive demo",
    "July 2026 · ~8 minutes end-to-end · ~$1/hr while running · companion to the full Setup & Benchmark Guide")

g.h1("1. What the automation does")
g.body("A single script (deploy/deploy.sh in the kvcache-demo repository) takes you from nothing to a "
       "browser dashboard driving a live A100: it rents the cheapest suitable GPU on vast.ai, pushes every "
       "component, installs NVIDIA NIXL, downloads the model, applies a required LMCache patch, starts the "
       "LMCache controller and the dashboard, launches vLLM in the full tiering hierarchy "
       "(GPU 47 GB → pinned CPU RAM 60 GB → NIXL/NVMe 85.6 GB), and opens an SSH tunnel so the dashboard "
       "appears at http://127.0.0.1:7811 on your machine. Everything is idempotent: re-running against the "
       "same instance repairs rather than duplicates.")

g.h1("2. Prerequisites (once per workstation)")
g.bullet("macOS or Linux with Python 3.10+ and an SSH client (Windows: use WSL).")
g.bullet("A vast.ai account with ~$5+ prepaid credit (cloud.vast.ai → Billing → Add Credit).")
g.bullet("Git, to clone the repository.")

g.h1("3. One-time setup")
g.h2("3.1 vast.ai CLI and API key")
g.code(["python3 -m pip install --user vastai",
        "# pip installs off-PATH; add it permanently:",
        "echo 'export PATH=\"$(python3 -m site --user-base)/bin:$PATH\"' >> ~/.zshrc && source ~/.zshrc",
        "vastai set api-key <KEY-from-cloud.vast.ai-Account-page>",
        "vastai show user --raw | grep credit   # sanity: your balance appears"])
g.h2("3.2 SSH key registered with vast.ai")
g.code(["ssh-keygen -t ed25519 -f ~/.ssh/id_ed25519 -N \"\" -C \"vast\"   # skip if the key exists",
        "vastai create ssh-key \"$(cat ~/.ssh/id_ed25519.pub)\""])
g.h2("3.3 Clone the repository")
g.code(["git clone https://github.com/cypanized/kvcache-demo.git ~/kvcache",
        "cd ~/kvcache"])

g.h1("4. Deploy")
g.code(["./deploy/deploy.sh                # rent cheapest suitable A100 and deploy",
        "# or:",
        "./deploy/deploy.sh <INSTANCE_ID>  # (re)deploy onto an instance you already rented"])
g.body("Expected output, stage by stage (typical timings on a fresh box):", after=4)
g.code(["== searching offers ==                      (~5 s — picks cheapest A100 matching the filter)",
        "== rented instance 45246393 ==              (billing starts here)",
        "== waiting for instance + SSH ==            (~2-4 min — docker image pull; auto-starts if",
        "                                             the host parks the fresh instance as 'exited')",
        "== pushing demo files ==                    (~5 s)",
        "== dispatching bootstrap (detached) ==      (survives SSH drops)",
        "== starting local tunnel ==                 http://127.0.0.1:7811",
        "== waiting for the demo to come up ==       (~3-5 min — model download + vLLM load)",
        "✔ DEMO READY — open http://127.0.0.1:7811"])
g.callout("WHAT THE HIDDEN BOOTSTRAP DOES ON THE BOX", "installs the nixl wheel into vLLM's venv · "
          "downloads Qwen2.5-14B-Instruct (~28 GB, skipped if cached) · generates the 40-doc corpus · "
          "patches LMCache 0.5.1's NIXL backend with a clear() method (upstream lacks one) · writes and "
          "starts the controller (:9050) and dashboard (:7811) launchers with PYTHONHASHSEED pinned · "
          "launches vLLM in full-hierarchy mode via a crash-proof detached switch script.")

g.h1("5. The rented hardware — what the filter guarantees")
g.code(["gpu_name in [A100_SXM4, A100_PCIE] · gpu_ram >= 75 GB · disk >= 150 GB",
        "reliability > 0.98 · inet_down >= 500 Mbps · verified hosts · CUDA >= 12.4"])
g.body("Typical price $0.85–1.30/hr. Host differences (cgroup v1 vs v2, exited-on-create, SSH proxy "
       "flakiness) are handled by the scripts — validated across multiple distinct hosts.")

g.h1("6. Using the demo")
g.bullet("the Architecture button (header) opens a diagram of every deployed piece and the request/KV/admin paths.",
         bold_lead="Start with the map: ")
g.bullet("load docs (2 GB KV each), run manual 10-doc sweeps, clear individual tiers with receipts, paste "
         "your own prompts; the monitor column shows per-request TTFT, decode, and totals; tier cards show "
         "live residency and open evidence terminals on click.",
         bold_lead="Then drive it: ")
g.bullet("the six-step script is on the page under 'Suggested experiment' — cold sweep → GPU hits → flood "
         "the pyramid → natural spill-over serving → per-tier clears → total amnesia.",
         bold_lead="Guided walkthrough: ")

g.h1("7. Troubleshooting a deploy")
g.code(["# demo never reported ready — inspect on the box:",
        "ssh -p <PORT> root@<HOST> 'tail -40 /root/bootstrap.log /root/server_full.log'",
        "# dashboard unreachable but instance fine — restart the local tunnel:",
        "kill $(cat /tmp/kvdemo_tunnel.pid) 2>/dev/null",
        "nohup bash -c 'while true; do ssh -N -L 7811:127.0.0.1:7811 -p <PORT> root@<HOST>; sleep 3; done' &",
        "# dashboard process died on the box:",
        "ssh -p <PORT> root@<HOST> 'setsid nohup bash /root/restart_ui.sh >/dev/null 2>&1 < /dev/null &'"])

g.h1("8. Tear down (billing stops immediately)")
g.code(["vastai destroy instance <ID>",
        "vastai show instances        # confirm empty",
        "kill $(cat /tmp/kvdemo_tunnel.pid) 2>/dev/null   # stop the local tunnel"])
g.callout("COST DISCIPLINE", "A stopped instance still bills for storage — destroy, don't stop. "
          "Everything is reproducible: the next ./deploy/deploy.sh rebuilds the identical demo in ~8 minutes.")

g.save("/Users/wongtran/kvcache/report/KVCache-Demo-Automated-Deployment.docx")


# ==================================================================
# GUIDE 2 — MANUAL DEPLOYMENT
# ==================================================================
m = GuideDoc(
    "KV Cache Tiering Demo — Manual Deployment",
    "Every component installed, configured, and verified by hand on a vast.ai A100",
    "July 2026 · ~45–60 minutes hands-on · what deploy.sh automates, unpacked step by step")

m.h1("1. The system you are building")
m.body("Six cooperating pieces on one GPU node, plus a tunnel from your workstation:")
m.image("/Users/wongtran/kvcache/report/architecture.png", 6.4,
        "vLLM + LMCache + NIXL and the three KV tiers (demo runs the full hierarchy: all tiers at once)")
m.bullet("serves Qwen2.5-14B on :8000; owns Tier 0 (GPU paged KV, ~47 GB).", bold_lead="vLLM — ")
m.bullet("in-process KV manager: 256-token chunks, content-hash keys, write-through to every enabled tier, per-tier LRU.", bold_lead="LMCache — ")
m.bullet("Tier 1: 60 GB pinned host RAM (LocalCPUBackend).", bold_lead="CPU tier — ")
m.bullet("Tier 2: NVIDIA NIXL agent + POSIX plugin → 1,700 × 48 MiB slot files = 85.6 GB on NVMe.", bold_lead="NIXL tier — ")
m.bullet("admin service on :9050 — pure per-tier eviction (/clear) and index audits (/lookup).", bold_lead="LMCache controller — ")
m.bullet("dashboard server on :7811 + probes (TTFT measurement, evidence, crash-proof mode switching).", bold_lead="Demo tooling — ")

m.h1("2. Rent and reach the instance")
m.code(["vastai search offers 'gpu_name in [A100_SXM4,A100_PCIE] num_gpus=1 gpu_ram>=75 \\",
        "  disk_space>=150 reliability>0.98 inet_down>=500 rentable=true verified=true \\",
        "  cuda_vers>=12.4' -o 'dph'",
        "vastai create instance <OFFER_ID> --image lmcache/vllm-openai:latest --disk 150 --ssh --direct",
        "# note new_contract = your INSTANCE_ID; poll until running, then:",
        "vastai show instance <ID> --raw | grep -E '\"ssh_host\"|\"ssh_port\"|actual_status'",
        "ssh -p <PORT> root@<HOST>    # type yes at the fingerprint prompt; nvidia-smi must show the A100"])
m.callout("HOST QUIRK #1", "Fresh instances sometimes boot to 'exited' after the image pull. "
          "Fix: vastai start instance <ID>, wait for running again.")

m.h1("3. Install the components")
m.h2("3.1 vLLM + LMCache")
m.body("The lmcache/vllm-openai:latest image ships both, version-matched, in /opt/venv (validated pair: "
       "vLLM 0.24.0 + LMCache 0.5.1). Nothing to install. On a bare CUDA machine instead:", after=4)
m.code(["uv venv /opt/venv --python 3.12",
        "uv pip install --python /opt/venv/bin/python3 vllm==0.24.0 lmcache==0.5.1 nixl"])
m.callout("PIN THE PAIR", "LMCache compiles against a specific vLLM connector API. Install both in one "
          "command with pinned versions — independent resolution produces installs that break at runtime, "
          "not install time.")
m.h2("3.2 NVIDIA NIXL (on the docker image)")
m.code(["# the image's venv has no pip module — use uv, which is preinstalled:",
        "uv pip install --python /opt/venv/bin/python3 nixl",
        "python3 -c 'import nixl; print(\"nixl OK\")'",
        "ls /opt/venv/lib/python3.12/site-packages/.nixl*/plugins/   # POSIX, GDS, GDS_MT, OBJ, UCX…"])
m.h2("3.3 Model and corpus")
m.code(["python3 -c \"from huggingface_hub import snapshot_download; snapshot_download('Qwen/Qwen2.5-14B-Instruct')\"",
        "# demo files: clone the repo on your workstation and push them up",
        "git clone https://github.com/cypanized/kvcache-demo.git ~/kvcache && cd ~/kvcache",
        "scp -P <PORT> remote/run_server.sh remote/gen_docs.py demo/switch.sh demo/ask.py \\",
        "  demo/spill_test.py demo/evidence_lookup.py demo/onbox_server.py demo/index.html root@<HOST>:/root/",
        "ssh -p <PORT> root@<HOST> 'chmod +x /root/*.sh && cd /root && python3 gen_docs.py 40 9000'"])
m.body("The 40 synthetic docs are ~65 KB of deterministic text each (2.6 MB total) but cost ~2 GB of KV "
       "each when read — 192 KiB per token × ~10,500 tokens. That 30,000× inflation is the entire reason "
       "tiering exists.")

m.h1("4. Patch LMCache 0.5.1 (required for per-tier clears)")
m.body("Upstream's NIXL backend implements no clear() — the controller's per-tier eviction silently "
       "reports 0 tokens. Add one (idempotent; bootstrap.sh does exactly this):", after=4)
m.code(["python3 - <<'EOF'",
        "path = \"/opt/venv/lib/python3.12/site-packages/lmcache/v1/storage_backend/nixl_storage_backend.py\"",
        "src = open(path).read()",
        "if \"def clear(self) -> int:\" not in src:",
        "    anchor = \"    async def mem_to_storage(\"",
        "    patch = '''    def clear(self) -> int:",
        "        with self.key_lock:",
        "            num_chunks = len(self.key_dict)",
        "            for meta in self.key_dict.values():",
        "                try: self.pool.push(meta.index)",
        "                except Exception: pass",
        "            self.key_dict.clear()",
        "        return num_chunks * 256",
        "",
        "'''",
        "    open(path, \"w\").write(src.replace(anchor, patch + anchor, 1))",
        "EOF"])

m.h1("5. Configure and launch the stack")
m.h2("5.1 The full-hierarchy configuration, line by line")
m.code(["# /root/lmcache_full.yaml  (run_server.sh writes this; shown for understanding)",
        "chunk_size: 256                  # tokens per cache chunk (= 48 MiB of KV each)",
        "local_cpu: true                  # enable Tier 1",
        "max_local_cpu_size: 60           # GB of pinned host RAM",
        "nixl_buffer_device: \"cpu\"        # POSIX plugin requires CPU staging (cuda → assert fail)",
        "enable_controller: true          # register with the admin controller",
        "lmcache_instance_id: \"demo\"",
        "controller_pull_url: \"localhost:8300\"",
        "controller_reply_url: \"localhost:8400\"",
        "lmcache_worker_ports: [8500]",
        "extra_config:",
        "  enable_nixl_storage: true      # Tier 2 via NIXL",
        "  nixl_backend: \"POSIX\"          # GDS/GDS_MT on GDS-capable hardware",
        "  nixl_pool_size: 1700           # slot COUNT (× 48 MiB = 85.6 GB) — NEVER bytes",
        "  nixl_path: \"/root/lmcache_nixl/\""])
m.callout("THREE LANDMINES IN THIS CONFIG", "(1) nixl_pool_size in bytes makes LMCache build a "
          "multi-hundred-GB index and the engine is OOM-killed with no traceback. (2) The pool keeps one "
          "open fd per slot — launch with ulimit -n 65536. (3) Do NOT add use_direct_io: true — it "
          "deadlocks 0.5.1 at startup (futex wait, no error).")
m.h2("5.2 Environment the launcher must set (run_server.sh does all three)")
m.code(["export VLLM_SERVER_DEV_MODE=1   # exposes /reset_prefix_cache (pure GPU-tier eviction)",
        "export PYTHONHASHSEED=0         # chunk keys use Python's randomized hash — every LMCache",
        "                                # process needs the same seed or controller lookups match nothing",
        "ulimit -n 65536                 # NIXL pool fds"])
m.h2("5.3 The controller service")
m.code(["cat > /root/start_controller.sh <<'EOF'",
        "#!/bin/bash",
        "export PYTHONHASHSEED=0",
        "pkill -f 'lmcache.v1.api_serve[r]' 2>/dev/null; sleep 1",
        "cd /root && exec python3 -m lmcache.v1.api_server --host 127.0.0.1 --port 9050 \\",
        "  --monitor-ports '{\"pull\": 8300, \"reply\": 8400}' > /root/controller.log 2>&1",
        "EOF",
        "chmod +x /root/start_controller.sh",
        "setsid nohup bash /root/start_controller.sh >/dev/null 2>&1 </dev/null &"])
m.h2("5.4 The dashboard server")
m.code(["cat > /root/restart_ui.sh <<'EOF'",
        "#!/bin/bash",
        "pkill -f 'onbox_serve[r].py' 2>/dev/null; sleep 2",
        "cd /root && exec python3 onbox_server.py > /root/onbox_server.log 2>&1",
        "EOF",
        "chmod +x /root/restart_ui.sh",
        "setsid nohup bash /root/restart_ui.sh >/dev/null 2>&1 </dev/null &"])
m.h2("5.5 Launch vLLM in full-hierarchy mode")
m.code(["setsid nohup /root/switch.sh full >/dev/null 2>&1 </dev/null &",
        "# switch.sh: kills any prior server INCLUDING orphaned 'VLLM::EngineCore' processes",
        "# (they rename themselves; a naive pkill leaves them holding ~45 GB of GPU memory),",
        "# wipes stale tier files, then launches run_server.sh full — detached, so an SSH",
        "# drop mid-switch cannot kill the restart. Load takes ~3 minutes."])
m.h2("5.6 Tunnel from your workstation")
m.code(["nohup bash -c 'while true; do ssh -o ServerAliveInterval=30 -o ExitOnForwardFailure=yes \\",
        "  -N -L 7811:127.0.0.1:7811 -p <PORT> root@<HOST>; sleep 3; done' >/dev/null 2>&1 &",
        "# then open http://127.0.0.1:7811"])

m.h1("6. Verify every layer (do not skip)")
m.code(["# 1. engine + backends healthy, nothing degraded:",
        "grep -cE 'init failed|degraded' /root/server_full.log            # MUST print 0",
        "grep -o 'Created backend: [A-Za-z]*' /root/server_full.log | sort -u",
        "#    → LocalCPUBackend + NixlStorageBackend",
        "# 2. the tier ladder — each rung via pure evictions (expected TTFTs):",
        "#    cold ~1.9s → GPU hit ~0.07s → clear GPU → CPU hit ~0.17s",
        "#    → clear GPU+CPU → NIXL hit ~0.35s → clear all → ~1.8s recompute",
        "# 3. which tier served? read the engine's own log — throughput fingerprints the medium:",
        "grep -a 'Retrieved' /root/server_full.log | tail -3",
        "#    ≈24 GB/s = pinned RAM (Tier 1) · ≈5 GB/s = NVMe pool (Tier 2)",
        "# 4. controller audit matches reality:",
        "python3 /root/evidence_lookup.py     # per-doc /lookup vs the dashboard gauges"])
m.callout("KNOWN 0.5.1 BLIND SPOT", "The NIXL backend never reports its contents to the controller "
          "registry, so /lookup cannot name NixlStorageBackend as a location — after clearing the CPU "
          "tier the audit reads empty while the pool still serves (prove it: clear GPU+CPU, ask a doc, "
          "~0.35s + a 'Retrieved 10496 of 10496' log line). LMCache issue #3387 tracks the same failure "
          "class for P2P.")

m.h1("7. Troubleshooting quick table")
m.code(["engine dies at start, exit 137, no traceback → nixl_pool_size was bytes; use a slot count",
        "'Too many open files' during pool creation   → ulimit -n 65536 before launch",
        "'Invalid NIXL backend & device combination'  → POSIX needs nixl_buffer_device: cpu",
        "warm TTFT equals baseline in a tiered mode   → LMCache degraded; grep 'init failed'",
        "startup hangs forever (futex wait)           → remove use_direct_io from extra_config",
        "next launch fails; GPU busy with no server   → pkill -9 -f 'VLLM::' (orphaned engine cores)",
        "controller /lookup or /clear returns 0        → PYTHONHASHSEED not pinned in ALL processes,",
        "                                               or tokens lack the chat template",
        "host resource stats error on some hosts      → cgroup v1 vs v2 paths (onbox_server handles both)"])

m.h1("8. Tear down")
m.code(["vastai destroy instance <ID>   # irreversible; billing stops; storage is NOT free on stopped instances"])

m.save("/Users/wongtran/kvcache/report/KVCache-Demo-Manual-Deployment.docx")
