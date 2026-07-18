#!/usr/bin/env python3
"""Generate KVCache-Tiering-Setup-Guide.docx — VAST-styled step-by-step guide."""
import copy

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

DEEP_BLUE = RGBColor(0x1E, 0x3A, 0x5F)
DEEP_BLUE_HEX = "1E3A5F"
LIGHT_BLUE_HEX = "F0F7FA"
GRAY_HEX = "E8E8E8"
CODE_BG = "F5F5F5"
BORDER_GRAY = "CCCCCC"

LOGO = ("/Users/wongtran/Library/Application Support/Claude/local-agent-mode-sessions/"
        "skills-plugin/3928ab7b-2926-439f-8a2e-ea3d7739d43a/cb212373-0fdf-4237-9256-3eefc9914e53/"
        "skills/vast-docx-style/assets/vast_logo.png")

doc = Document()

# --- base styles ---
normal = doc.styles["Normal"]
normal.font.name = "Arial"
normal.font.size = Pt(11)
for lvl, sz, before, after in (("Heading 1", 18, 14, 7), ("Heading 2", 14, 11, 5), ("Heading 3", 12, 9, 4)):
    st = doc.styles[lvl]
    st.font.name = "Arial"
    st.font.size = Pt(sz)
    st.font.bold = True
    st.font.color.rgb = DEEP_BLUE
    st.paragraph_format.space_before = Pt(before)
    st.paragraph_format.space_after = Pt(after)
    # kill the default heading underline/border inheritance from theme fonts
    st.element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")

sec = doc.sections[0]
sec.page_width, sec.page_height = Inches(8.5), Inches(11)
sec.left_margin = sec.right_margin = sec.top_margin = sec.bottom_margin = Inches(1)

# header with logo
hdr_p = sec.header.paragraphs[0]
hdr_p.paragraph_format.space_after = Pt(0)
run = hdr_p.add_run()
run.add_picture(LOGO, width=Inches(0.52))


def shade(el, hex_fill):
    shd = el.makeelement(qn("w:shd"), {qn("w:val"): "clear", qn("w:fill"): hex_fill})
    el.get_or_add_pPr().append(shd) if el.tag.endswith("}p") else el.append(shd)


def body(text, bold_lead=None, after=6, italic=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(after)
    if bold_lead:
        r = p.add_run(bold_lead)
        r.bold = True
    r = p.add_run(text)
    r.italic = italic
    return p


def bullet(text, bold_lead=None):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(3)
    if bold_lead:
        r = p.add_run(bold_lead)
        r.bold = True
    p.add_run(text)
    return p


def code(lines, after=8):
    if isinstance(lines, str):
        lines = lines.split("\n")
    for i, line in enumerate(lines):
        p = doc.add_paragraph()
        pPr = p._p.get_or_add_pPr()
        pPr.append(p._p.makeelement(qn("w:shd"), {qn("w:val"): "clear", qn("w:fill"): CODE_BG}))
        p.paragraph_format.space_after = Pt(after if i == len(lines) - 1 else 0)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.left_indent = Inches(0.12)
        r = p.add_run(line if line else " ")
        r.font.name = "Consolas"
        r._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")
        r._element.rPr.rFonts.set(qn("w:cs"), "Consolas")
        r.font.size = Pt(9)


def callout(label, text):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pPr.append(p._p.makeelement(qn("w:shd"), {qn("w:val"): "clear", qn("w:fill"): LIGHT_BLUE_HEX}))
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.left_indent = Inches(0.12)
    r = p.add_run(label + "  ")
    r.bold = True
    r.font.color.rgb = DEEP_BLUE
    p.add_run(text)


def make_table(headers, rows, widths, highlight_alt=True):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.style = "Table Grid"
    # borders to light gray
    tbl_pr = t._tbl.tblPr
    borders = t._tbl.makeelement(qn("w:tblBorders"), {})
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        e = t._tbl.makeelement(qn(f"w:{edge}"), {qn("w:val"): "single", qn("w:sz"): "4", qn("w:color"): BORDER_GRAY})
        borders.append(e)
    tbl_pr.append(borders)
    for j, (h, w) in enumerate(zip(headers, widths)):
        c = t.rows[0].cells[j]
        c.width = Inches(w)
        c._tc.get_or_add_tcPr().append(t._tbl.makeelement(qn("w:shd"), {qn("w:val"): "clear", qn("w:fill"): DEEP_BLUE_HEX}))
        p = c.paragraphs[0]
        r = p.add_run(h)
        r.bold = True
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        r.font.size = Pt(10)
    for i, row in enumerate(rows):
        for j, (val, w) in enumerate(zip(row, widths)):
            c = t.rows[i + 1].cells[j]
            c.width = Inches(w)
            if highlight_alt and i % 2 == 0:
                c._tc.get_or_add_tcPr().append(t._tbl.makeelement(qn("w:shd"), {qn("w:val"): "clear", qn("w:fill"): LIGHT_BLUE_HEX}))
            p = c.paragraphs[0]
            r = p.add_run(str(val))
            r.font.size = Pt(10)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return t


# ============ TITLE ============
t = doc.add_paragraph()
t.paragraph_format.space_before = Pt(24)
r = t.add_run("KV Cache Tiering with LMCache and NVIDIA NIXL")
r.font.size = Pt(24)
r.bold = True
r.font.color.rgb = DEEP_BLUE
s = doc.add_paragraph()
r = s.add_run("Step-by-step setup and benchmark guide — vLLM on a rented vast.ai A100")
r.font.size = Pt(13)
r.font.color.rgb = RGBColor(0x55, 0x66, 0x77)
meta = doc.add_paragraph()
r = meta.add_run("Version 2.0 · July 2026 · Benchmark + interactive demo, validated end-to-end on vast.ai A100 instances · Manual path ≈ 2 h or one-command deploy ≈ 8 min (section 12.7)")
r.font.size = Pt(9)
r.italic = True
r.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

# ============ 1. OVERVIEW ============
doc.add_heading("1. What you will build", level=1)
body("You will rent a single A100 80GB GPU on vast.ai, serve Qwen2.5-14B-Instruct with vLLM, and "
     "benchmark four key-value (KV) cache configurations against an identical workload. The workload is "
     "deliberately sized so its KV footprint (~290k tokens) overflows the GPU cache budget (~75k tokens) "
     "by 4x — the failure mode that makes tiering matter. Without tiering, every repeat request re-runs "
     "a full 12k-token prefill; with LMCache tiering, repeat requests fetch previously computed KV blocks "
     "from CPU RAM or NVMe storage through NVIDIA NIXL.")
body("Results you should be able to reproduce (time-to-first-token on repeat requests):", after=4)
make_table(
    ["Configuration", "KV path", "Warm TTFT", "Speedup"],
    [["Baseline vLLM", "GPU HBM only (prefix cache thrashed)", "1.741 s", "1.0x"],
     ["LMCache · CPU", "GPU → pinned CPU RAM (80 GB)", "0.170 s", "10.3x"],
     ["LMCache · Disk", "GPU → CPU (20 GB) → NVMe files", "0.434 s", "4.0x"],
     ["LMCache · NIXL", "GPU → NIXL POSIX plugin → NVMe pool", "0.362 s", "4.8x"]],
    [1.5, 3.2, 0.9, 0.9])
callout("WHY IT MATTERS", "The NVMe file pool stands in for any NIXL-reachable external target. In "
        "production, the same wiring points at a VAST NFS or GDS mount — KV capacity decoupled from the "
        "GPU node, surviving restarts, shared across a fleet.")
doc.add_heading("1.1 Terms used in this guide", level=2)
make_table(
    ["Term", "Meaning"],
    [["KV cache", "The attention key/value tensors an LLM computes for every token it has read. Keeping them means never re-reading the same text."],
     ["Prefill", "The expensive first phase of a request: the model reads the whole prompt and builds its KV cache before generating anything."],
     ["TTFT", "Time to first token — how long the user waits before the first word of the answer appears. Prefill dominates it."],
     ["Cold / warm", "Cold: the prompt has never been seen (full prefill). Warm: its KV blocks are already cached somewhere."],
     ["Tier", "A place KV blocks can live, fastest first: GPU memory (HBM), CPU RAM, then external storage (NVMe / network)."],
     ["Instance", "The rented cloud machine on vast.ai. You control it from your own computer through SSH."],
     ["SSH / scp", "Secure shell — a remote command line on the instance — and its file-copy companion."]],
    [1.2, 5.3])

# ============ 2. ARCHITECTURE ============
doc.add_heading("2. How the pieces fit together", level=1)
body("Three independent projects cooperate, joined by two narrow interfaces. vLLM owns inference and its "
     "GPU paged KV cache; a KV connector hook hands finished KV blocks to LMCache. LMCache owns the cache "
     "hierarchy — chunking token sequences into 256-token blocks, hashing them into cache keys, and "
     "deciding which tier holds each block. When a block must leave the node's RAM, LMCache's NIXL storage "
     "backend hands the transfer to NVIDIA NIXL, whose pluggable backends (POSIX, GDS, S3-style OBJ) move "
     "bytes to the actual storage target.")
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.add_run().add_picture("/Users/wongtran/kvcache/report/architecture.png", width=Inches(6.4))
cap = doc.add_paragraph()
cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = cap.add_run("Figure 1 — vLLM, LMCache, and NIXL: the KV store/retrieve paths and the three cache tiers")
r.italic = True
r.font.size = Pt(9)
r.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
body("Reading the diagram top to bottom:", after=4)
bullet("On a cache miss, vLLM prefills normally; the connector then pushes the new "
       "KV blocks down — LMCache keeps them in its pinned CPU pool and/or streams them "
       "through NIXL to the storage pool.", bold_lead="Store path (downward): ")
bullet("On a later request with the same prefix, LMCache finds the chunks by hash, pulls "
       "them up from the fastest tier that has them, and vLLM skips prefill for those tokens — that skip "
       "is the entire speedup measured in this guide.", bold_lead="Retrieve path (upward): ")
bullet("Each mode in section 8 simply enables deeper tiers: baseline stops at tier 0 (GPU), "
       "cpu adds tier 1 (local_cpu), disk and nixl add tier 2 (NVMe — via LMCache's own file backend or "
       "via NIXL respectively).", bold_lead="Mapping to the four configs: ")
bullet("The POSIX plugin targets any filesystem path. Point nixl_path at a VAST NFS or "
       "GDS mount instead of local NVMe and nothing else changes — that is the production topology, with "
       "cache capacity decoupled from the GPU node and shared across a fleet.", bold_lead="Swapping the bottom box: ")

# ============ 3. BEFORE YOU START ============
doc.add_heading("3. Before you start", level=1)
body("Time and money map — every paid minute is on the rented GPU, so the meter only runs from section 6 "
     "to section 12:", after=4)
make_table(
    ["Stage", "Sections", "Time", "Cost"],
    [["Account, terminal, CLI setup", "3–5", "~20 min", "$0"],
     ["Rent GPU + prepare instance", "6–7", "~25 min", "~$0.50"],
     ["Configure, verify, benchmark 4 modes", "8–9", "~50 min", "~$1.00"],
     ["Analyze + tear down", "10, 13", "~15 min", "~$0.25"]],
    [2.6, 1.2, 1.2, 1.2])
doc.add_heading("3.1 Create a vast.ai account and add credit", level=2)
bullet("Go to cloud.vast.ai in a browser and sign up (email + password, or Google).")
bullet("Top-right menu → Billing → Add Credit. Add at least $5 by card. vast.ai is prepaid — rentals "
       "draw down this balance, and nothing can charge beyond it.")
bullet("Stay logged in; you will come back for the API key in section 4.")
doc.add_heading("3.2 Open a terminal", level=2)
body("Every command in this guide is typed into a terminal. On macOS open Applications → Utilities → "
     "Terminal (or search \"Terminal\" in Spotlight). On Linux, any terminal emulator. On Windows, "
     "install WSL (Ubuntu) first and run everything inside it — the commands here assume a Unix shell. "
     "Commands shown in gray boxes are typed at the prompt; press Enter after each line. Lines starting "
     "with # are comments — don't type them.", after=8)
doc.add_heading("3.3 Check Python", level=2)
code(["python3 --version   # want 3.10 or newer"])
body("If that prints \"command not found\": on macOS install the Xcode command line tools with "
     "xcode-select --install, or download Python from python.org; on Ubuntu/WSL run "
     "sudo apt install python3-pip.", after=8)
body("You will also need: nothing else. No HuggingFace token (the model is ungated), no GPU on your own "
     "machine, no docker locally. Validated total cost of the original run: about $2.50.", after=8)

# ============ 4. CLI ============
doc.add_heading("4. Install the vast.ai CLI and add your API key", level=1)
body("The vastai command lets you search, rent, and destroy GPU machines from the terminal instead of "
     "clicking through the website. Install it with Python's package manager:", after=4)
code(["python3 -m pip install --user vastai"])
body("pip installs the command into a per-user folder that is usually NOT on your PATH (the list of "
     "folders the shell searches for commands). Find it and add it permanently:", after=4)
code(["python3 -m site --user-base    # prints e.g. /Users/you/Library/Python/3.11",
      "# append /bin to that path and add it to your shell profile:",
      "echo 'export PATH=\"$(python3 -m site --user-base)/bin:$PATH\"' >> ~/.zshrc",
      "source ~/.zshrc                # bash users: use ~/.bashrc in both lines",
      "vastai --help                  # should print the command list, not 'command not found'"])
body("Now connect the CLI to your account. In the browser: cloud.vast.ai → click your avatar (top "
     "right) → Account → API Keys → copy the key. Then:", after=4)
code(["vastai set api-key <PASTE_YOUR_KEY_HERE>",
      "vastai show user --raw | grep credit    # sanity check: prints the balance you added"])

doc.add_heading("5. Register an SSH key", level=1)
body("SSH is how you will control the rented machine. It authenticates with a key pair: a private file "
     "that never leaves your computer, and a public file you give to vast.ai, which installs it on any "
     "instance you rent. Create the pair (skip the first command if ~/.ssh/id_ed25519 already exists) "
     "and upload the public half:", after=4)
code(["ssh-keygen -t ed25519 -f ~/.ssh/id_ed25519 -N \"\" -C \"vast-kvcache\"",
      "vastai create ssh-key \"$(cat ~/.ssh/id_ed25519.pub)\"   # expect: {'success': True, ...}"])

# ============ 6. RENT ============
doc.add_heading("6. Find and rent the GPU", level=1)
doc.add_heading("6.1 Search the market", level=2)
body("vast.ai is a marketplace, so prices and machines vary by the hour. Search for a single A100 80GB "
     "with fast networking, plenty of host RAM (the CPU tier needs it), and 150+ GB of disk (the model "
     "is 28 GB; the KV pools want 60–120 GB). The billing meter does NOT start at search — only at "
     "create instance in 6.2:", after=4)
code(["vastai search offers 'gpu_name in [A100_SXM4,A100_PCIE] num_gpus=1 gpu_ram>=75 \\",
      "  disk_space>=150 reliability>0.98 inet_down>=500 rentable=true verified=true \\",
      "  cuda_vers>=12.4' -o 'dph'"])
body("The result is a table sorted by price. The number you need is the ID column (a long number like "
     "34227706) — not the # row counter. Prefer a row with RAM above 200, Net_down above 1000 (that is "
     "Mbit/s — faster model download), R (reliability) above 99, and $/hr near 1.00.", after=8)
doc.add_heading("6.2 Rent it (billing starts here)", level=2)
body("Rent the chosen offer with the LMCache docker image, which ships vLLM and LMCache pre-installed "
     "and version-matched:", after=4)
code(["vastai create instance <OFFER_ID> --image lmcache/vllm-openai:latest \\",
      "  --disk 150 --ssh --direct --label kvcache-demo",
      "# reply: {'success': True, 'new_contract': 45214089, ...}",
      "#                          ^^^^^^^^^^^^^^^^^^^^^^^^",
      "#         this number is your INSTANCE_ID — note it down"])
doc.add_heading("6.3 Wait for it to boot, then get the address", level=2)
body("First boot takes 3–6 minutes while the ~15 GB image downloads. Re-run this until status says "
     "running, then note the host and port:", after=4)
code(["vastai show instance <INSTANCE_ID> --raw | grep -E '\"ssh_host\"|\"ssh_port\"|actual_status'",
      "#   \"actual_status\": \"running\",",
      "#   \"ssh_host\": \"ssh6.vast.ai\",     <- your HOST",
      "#   \"ssh_port\": 14088,               <- your PORT"])
doc.add_heading("6.4 Connect for the first time", level=2)
body("Everywhere below, replace <HOST> and <PORT> with the two values from 6.3:", after=4)
code(["ssh -p <PORT> root@<HOST>",
      "# first connection asks: 'Are you sure you want to continue connecting?' — type yes",
      "nvidia-smi     # run this on the instance: it must list one NVIDIA A100 80GB",
      "exit           # come back to your own machine"])
callout("CHECKPOINT", "You saw the A100 in nvidia-smi output. If ssh says 'Permission denied', wait 30 "
        "seconds and retry (the key may still be propagating), and confirm section 5 succeeded. From here "
        "on, the instance bills ~$1/hr whether you use it or not — if you need to pause for a day, skip "
        "ahead to section 12 and destroy it, then re-do sections 6–7 later.")

# ============ 7. PREP ============
doc.add_heading("7. Prepare the instance", level=1)
doc.add_heading("7.1 Copy the scripts to the instance", level=2)
body("The four helper scripts this guide uses are printed in full in Appendix A. Save each one on your "
     "own machine first: copy the text into a plain-text editor (TextEdit in plain-text mode, VS Code, "
     "nano) and save with the exact filename — gen_docs.py, benchmark.py, run_server.sh, run_mode.sh — "
     "all in one folder, e.g. ~/kvcache. Then push them to the instance and mark the shell scripts "
     "executable (note: scp uses a capital -P for the port):", after=4)
code(["cd ~/kvcache",
      "scp -P <PORT> gen_docs.py benchmark.py run_server.sh run_mode.sh root@<HOST>:/root/",
      "ssh -p <PORT> root@<HOST> 'chmod +x /root/run_server.sh /root/run_mode.sh && ls -la /root/'"])
body("Commands in the rest of section 7 run ON THE INSTANCE — connect with ssh -p <PORT> root@<HOST> "
     "first. You are root there; no sudo needed.", after=8)
doc.add_heading("7.2 Install NIXL into the serving venv", level=2)
body("The image's Python lives in /opt/venv but has no pip module; use uv, which is preinstalled:", after=4)
code(["uv pip install --python /opt/venv/bin/python3 nixl",
      "python3 -c 'import nixl; print(\"nixl OK\")'",
      "# optional: confirm the storage plugins shipped with the wheel",
      "ls /opt/venv/lib/python3.12/site-packages/.nixl*/plugins/ | grep -E 'POSIX|GDS'"])
doc.add_heading("7.3 Pre-download the model", level=2)
code(["python3 -c \"from huggingface_hub import snapshot_download; \\",
      "  snapshot_download('Qwen/Qwen2.5-14B-Instruct')\"",
      "du -sh /root/.cache/huggingface/hub/*Qwen*   # expect ~28G"])
doc.add_heading("7.4 Generate the benchmark corpus", level=2)
body("Still on the instance, create the 24 synthetic documents the benchmark reads:", after=4)
code(["cd /root && python3 gen_docs.py 24 9000   # 24 unique docs, ~12k tokens each",
      "# expect: wrote 24 docs, ~9000 words each -> docs.json"])
doc.add_heading("7.5 Where vLLM and LMCache come from (and installing them manually)", level=2)
body("This guide does not pip-install vLLM or LMCache because the docker image selected in section 6 — "
     "lmcache/vllm-openai:latest — ships both pre-installed and version-matched in /opt/venv (the "
     "validated run used vLLM 0.24.0 and LMCache 0.5.1). Using the image is strongly recommended: LMCache "
     "compiles native ops against a specific vLLM/PyTorch pair, and mismatched versions are the most "
     "common source of broken installs.", after=4)
body("If you are instead starting from a bare CUDA machine (your own server, or a plain "
     "nvidia/cuda or pytorch image), install the stack into a fresh virtualenv like this:", after=4)
code(["# Ubuntu 22.04+, NVIDIA driver >= 550 (CUDA 12.4+), Python 3.10-3.12",
      "curl -LsSf https://astral.sh/uv/install.sh | sh   # fast installer (or use python -m venv + pip)",
      "uv venv /opt/venv --python 3.12",
      "uv pip install --python /opt/venv/bin/python3 vllm==0.24.0 lmcache==0.5.1 nixl",
      "# verify all three import cleanly:",
      "/opt/venv/bin/python3 -c 'import vllm, lmcache, nixl; \\",
      "  print(vllm.__version__, lmcache.__version__, \"nixl OK\")'"])
callout("VERSION PINNING", "Install vllm and lmcache in one command with pinned, known-compatible "
        "versions (the pair above is what this guide validated). Letting pip resolve them separately can "
        "pull a newer vLLM whose connector API LMCache's release doesn't support yet — the failure shows "
        "up at runtime, not install time. If you want newer versions, take whatever pair a recent "
        "lmcache/vllm-openai image ships rather than resolving your own.")

# ============ 7. CONFIGS ============
doc.add_heading("8. The four server configurations", level=1)
body("All four use the same vLLM base command; the LMCache variants add a KV-transfer connector and a "
     "YAML config. The full launch script (run_server.sh) is in Appendix A — the essence of each mode:")
body("Two ways to use these configurations: the BENCHMARK path (this section + sections 9-10) runs the "
     "four single-tier modes to isolate each tier's numbers, while the INTERACTIVE DEMO (section 12) "
     "uses just two modes — baseline and the full hierarchy of 8.5 — because its per-tier clear buttons "
     "provide the isolation instead.", after=6)
body("The entire vLLM-to-LMCache integration is exactly two knobs on the vllm serve command — there is "
     "no other glue code:", after=4)
bullet("registers LMCache as vLLM's KV connector. kv_role: kv_both means this one process both "
       "saves KV after prefill and loads KV before prefill (the roles split only in "
       "disaggregated prefill/decode setups).",
       bold_lead="--kv-transfer-config '{\"kv_connector\":\"LMCacheConnectorV1\",\"kv_role\":\"kv_both\"}' — ")
bullet("an environment variable pointing at the YAML that tells LMCache which tiers to enable "
       "(CPU, disk, NIXL) and how big each is. Change the YAML, restart, and you have a different "
       "topology — vLLM flags never change.",
       bold_lead="LMCACHE_CONFIG_FILE=/root/lmcache_<mode>.yaml — ")
body("LMCache-to-NIXL wiring is in turn just the extra_config block inside that YAML (section 8.4): "
     "enable_nixl_storage picks the backend, nixl_backend picks the plugin, nixl_path picks the target. "
     "No NIXL daemon or separate service runs — the NIXL agent lives inside the vLLM server process.", after=8)
doc.add_heading("8.1 Baseline — GPU only", level=2)
code(["ulimit -n 65536",
      "vllm serve Qwen/Qwen2.5-14B-Instruct --host 127.0.0.1 --port 8000 \\",
      "  --gpu-memory-utilization 0.55 --max-model-len 16384"])
callout("GPU MEMORY UTILIZATION — 0.55 vs 0.92", "run_server.sh (Appendix A) defaults to 0.92, the honest "
        "setting: 80 GB minus 28 GB of weights leaves a ~47 GB / ~240k-token KV budget and the GPU is "
        "actually used. The benchmark table in section 1 was measured with GPU_UTIL=0.55 (a deliberate "
        "~15 GB / 75k-token cap) so the small 24-doc corpus would overflow the GPU 4x and thrash. To "
        "reproduce those numbers exactly, prefix each launch with GPU_UTIL=0.55; at 0.92 the working set "
        "must exceed ~47 GB (24+ docs) before the baseline starts thrashing. The setting is identical "
        "across configs either way — only the tiering changes.")
doc.add_heading("8.2 LMCache CPU tier", level=2)
code(["cat > /root/lmcache_cpu.yaml <<EOF",
      "chunk_size: 256",
      "local_cpu: true",
      "max_local_cpu_size: 80",
      "EOF",
      "LMCACHE_CONFIG_FILE=/root/lmcache_cpu.yaml vllm serve Qwen/Qwen2.5-14B-Instruct \\",
      "  --host 127.0.0.1 --port 8000 --gpu-memory-utilization 0.55 --max-model-len 16384 \\",
      "  --kv-transfer-config '{\"kv_connector\":\"LMCacheConnectorV1\",\"kv_role\":\"kv_both\"}'"])
doc.add_heading("8.3 LMCache disk tier", level=2)
code(["cat > /root/lmcache_disk.yaml <<EOF",
      "chunk_size: 256",
      "local_cpu: true",
      "max_local_cpu_size: 20",
      "local_disk: \"file:///root/lmcache_disk/\"",
      "max_local_disk_size: 120",
      "EOF",
      "# launch identical to 8.2 but with LMCACHE_CONFIG_FILE=/root/lmcache_disk.yaml"])
body("The CPU tier is deliberately squeezed to 20 GB (~14 of 24 documents) so a meaningful share of hits "
     "must come off NVMe rather than RAM.", after=8)
doc.add_heading("8.4 LMCache NIXL tier (the headline)", level=2)
code(["mkdir -p /root/lmcache_nixl",
      "cat > /root/lmcache_nixl.yaml <<EOF",
      "chunk_size: 256",
      "local_cpu: false",
      "max_local_cpu_size: 10",
      "nixl_buffer_device: \"cpu\"",
      "extra_config:",
      "  enable_nixl_storage: true",
      "  nixl_backend: \"POSIX\"",
      "  nixl_pool_size: 1500",
      "  nixl_path: \"/root/lmcache_nixl/\"",
      "EOF",
      "# launch identical to 8.2 but with LMCACHE_CONFIG_FILE=/root/lmcache_nixl.yaml"])
callout("CRITICAL — THE THREE NIXL GOTCHAS", "(1) nixl_pool_size is a descriptor COUNT, not bytes: one "
        "~48 MB chunk file per slot for this model (1,500 slots ≈ 72 GB). A bytes-sized value makes LMCache "
        "build a multi-hundred-GB index and the engine gets OOM-killed with no Python traceback. "
        "(2) The POSIX backend requires nixl_buffer_device: cpu — only GDS/GDS_MT/OBJ may stage in GPU "
        "memory; the wrong combination fails an assertion and LMCache silently degrades to recompute. "
        "(3) The file pool holds one open fd per slot — run ulimit -n 65536 before launching.")
doc.add_heading("8.5 The full-hierarchy mode (production topology)", level=2)
body("The four single-tier modes above isolate variables for benchmarking. run_server.sh also ships a "
     "fifth mode, full, which runs the complete pyramid in one server — the topology a production "
     "deployment actually uses: GPU ~47 GB (hardware) → pinned CPU RAM 60 GB (config) → NIXL pool "
     "1,700 slots × 48 MiB = 85.6 GB on NVMe. KV is written through to every tier; each tier "
     "LRU-evicts at its own cap while lower tiers retain their copies:", after=4)
code(["cat > /root/lmcache_full.yaml <<EOF",
      "chunk_size: 256",
      "local_cpu: true",
      "max_local_cpu_size: 60",
      "nixl_buffer_device: \"cpu\"",
      "enable_controller: true",
      "lmcache_instance_id: \"demo\"",
      "controller_pull_url: \"localhost:8300\"",
      "controller_reply_url: \"localhost:8400\"",
      "lmcache_worker_ports: [8500]",
      "extra_config:",
      "  enable_nixl_storage: true",
      "  nixl_backend: \"POSIX\"",
      "  nixl_pool_size: 1700",
      "  nixl_path: \"/root/lmcache_nixl/\"",
      "EOF",
      "/root/run_server.sh full   # or run_mode.sh full to also benchmark it"])
body("Verified tier ladder in this mode (each rung = clear the tiers above it, re-ask the same doc): "
     "cold prefill 1.95 s → GPU hit 0.07 s → CPU RAM 0.17 s → NIXL/NVMe 0.35 s → all tiers cleared "
     "1.81 s (recompute). Prefill is avoided as long as ANY tier still holds the KV.", after=8)
doc.add_heading("8.6 Verify the three layers are actually talking", level=2)
body("LMCache is designed to fail soft: if its backend cannot initialize, vLLM still serves correct "
     "answers at baseline speed. So after launching any LMCache mode, prove the wiring before "
     "benchmarking. First, the startup log (server_<mode>.log):", after=4)
code(["grep -E 'Initializing latest dev LMCache connector' server_nixl.log   # vLLM loaded the connector",
      "grep -E 'LMCache initialized for role' server_nixl.log               # LMCache engine is up",
      "grep -E 'Backend POSIX was instantiated' server_nixl.log             # NIXL plugin loaded (nixl mode)",
      "grep -cE 'init failed|degraded' server_nixl.log                      # MUST print 0"])
body("Second, a two-request smoke test — send the same long prompt twice and compare timings:", after=4)
code(["PROMPT=$(python3 -c \"print('word '*8000)\")",
      "time curl -s http://127.0.0.1:8000/v1/chat/completions -H 'Content-Type: application/json' \\",
      "  -d \"{\\\"model\\\":\\\"Qwen/Qwen2.5-14B-Instruct\\\",\\\"max_tokens\\\":8,\\",
      "       \\\"messages\\\":[{\\\"role\\\":\\\"user\\\",\\\"content\\\":\\\"$PROMPT\\\"}]}\" > /dev/null",
      "# run the identical curl again: the second call should be several times faster",
      "ls /root/lmcache_nixl/ | head -3 && du -sh /root/lmcache_nixl   # chunk files growing = KV landing"])
callout("PASS CRITERIA", "Second request markedly faster than the first, zero 'init failed' lines, and "
        "the tier's backing store growing (du on the nixl/disk path, or host RAM use for the CPU tier). "
        "If any of these fail, fix before running the benchmark — otherwise you will benchmark the "
        "baseline four times.")
doc.add_heading("8.7 Auditing what each tier actually holds (controller /lookup)", level=2)
body("Beyond timing signatures, LMCache can be asked directly which tier holds a given prompt's KV. "
     "Run the controller service (an admin API that LMCache workers register with), then query its "
     "/lookup endpoint with a prompt's token IDs — it answers with the location and token count from "
     "LMCache's own index. This is the definitive, read-only way to prove tier contents:", after=4)
code(["# one-time: run the controller and add controller keys to the LMCache YAML",
      "#   engine yaml: enable_controller: true, lmcache_instance_id, controller_pull_url,",
      "#   controller_reply_url, lmcache_worker_ports",
      "python3 -m lmcache.v1.api_server --host 127.0.0.1 --port 9050 \\",
      "  --monitor-ports '{\"pull\": 8300, \"reply\": 8400}' &",
      "",
      "# audit: tokenize the EXACT prompt, then look it up",
      "TOKENS=$(curl -s http://127.0.0.1:8000/tokenize -H 'Content-Type: application/json' \\",
      "  -d '{\"model\":\"Qwen/Qwen2.5-14B-Instruct\",\"add_generation_prompt\":true,",
      "       \"messages\":[{\"role\":\"user\",\"content\":\"<your prompt>\"}]}' | jq .tokens)",
      "curl -s http://127.0.0.1:9050/lookup -H 'Content-Type: application/json' \\",
      "  -d \"{\\\"tokens\\\": $TOKENS}\"",
      "# -> {\"layout_info\": {\"demo\": [\"LocalCPUBackend\", 10496]}}"])
callout("TWO REQUIREMENTS OR LOOKUPS SILENTLY RETURN NOTHING", "(1) Tokenize with the CHAT TEMPLATE "
        "(messages + add_generation_prompt), not the raw prompt text — the cache is keyed on the "
        "templated token sequence, so raw-text tokens never match. (2) Export PYTHONHASHSEED=0 (any "
        "fixed value) in EVERY LMCache process — the vLLM server AND the controller. LMCache's chunk "
        "keys use Python's builtin hash, which is randomized per process by default, so without a "
        "shared seed the controller can never reproduce the worker's keys. LMCache warns about this "
        "at startup; in a multi-process or multi-node deployment it is mandatory. The same two rules "
        "apply to the controller's /clear endpoint when targeting tokens rather than whole tiers.")
body("The controller also provides per-tier /clear (pure eviction with a dropped-token receipt) — "
     "location values are the backend names: LocalCPUBackend, LocalDiskBackend, NixlStorageBackend. "
     "Note: LMCache 0.5.1's NIXL backend does not implement clear() upstream (the storage manager "
     "skips it with a warning and reports 0 tokens); it needs a small patch or a newer release.", after=8)
callout("NIXL REGISTRY BLIND SPOT (0.5.1)", "The NIXL storage backend also never reports its contents "
        "to the controller registry (it has no controller message sender, unlike the CPU backend). "
        "Consequence: /lookup can never return NixlStorageBackend as a location — docs held only in the "
        "pool are invisible to the audit, and clearing the CPU tier makes lookups go empty even though "
        "the pool still serves every doc. Verify pool contents empirically instead: clear the GPU and "
        "CPU tiers, then re-ask a cached prompt — a ~0.35s TTFT plus an engine log line like 'Retrieved "
        "10496 out of 10496 required tokens' proves the pool answered.")

# ============ 8. RUN ============
doc.add_heading("9. Run the benchmark", level=1)
body("Each round runs unattended for ~10 minutes over SSH — and if your laptop sleeps or WiFi blips, a "
     "plain SSH session dies and kills the run with it. Protect against that with tmux, a program (already "
     "on the instance) that keeps a terminal session alive on the server even when you disconnect. Three "
     "commands are all you need:", after=4)
code(["ssh -p <PORT> root@<HOST>",
      "tmux new -s bench     # start (or: tmux attach -t bench to rejoin later)",
      "# ...run the benchmark commands below inside this session...",
      "# detach and leave it running: press Ctrl-b, release, then press d"])
body("Inside the tmux session, cycle the four configs with run_mode.sh — it kills the previous server, "
     "launches the new mode, waits for readiness, and runs 72 requests (24 docs x 3 passes):", after=4)
code(["/root/run_mode.sh baseline", "/root/run_mode.sh cpu",
      "rm -rf /root/lmcache_disk/*   # reclaim ~53 GB before the NIXL round",
      "/root/run_mode.sh disk", "/root/run_mode.sh nixl"])
body("Each round takes ~8–10 minutes (2–3 min model load + ~6 min of requests) and writes "
     "results_<mode>.json plus a cold/warm summary to stdout.", after=4)
callout("CHECKPOINT", "In the cpu/disk/nixl rounds, TTFT must drop sharply at request 25 (start of pass "
        "2). If warm numbers equal the baseline (~1.7 s), LMCache is in degraded mode — check the server "
        "log for 'LMCacheEngine marked as init failed'. Also verify KV is actually landing on storage: "
        "du -sh /root/lmcache_nixl should reach ~60–70 GB during the NIXL round.")

# ============ 9. ANALYZE ============
doc.add_heading("10. Analyze", level=1)
body("Back on your own machine (not the instance): save analyze.py from Appendix A.5 into your ~/kvcache "
     "folder, pull down the four result files, and summarize them:", after=4)
code(["cd ~/kvcache && mkdir -p results",
      "scp -P <PORT> 'root@<HOST>:/root/results_*.json' ./results/",
      "python3 analyze.py results/results_baseline.json results/results_cpu.json \\",
      "  results/results_disk.json results/results_nixl.json"])
body("Expected output (your absolute numbers will vary ±10% with host/NVMe speed; the ratios should not):", after=4)
code(["config        cold TTFT  warm TTFT   warm p95  vs baseline",
      "baseline         1.731s     1.741s     1.755s        1.00x",
      "cpu              1.829s     0.170s     0.180s       10.27x",
      "disk             1.854s     0.434s     0.513s        4.01x",
      "nixl             1.831s     0.362s     0.417s        4.81x"])

# ============ 10. TROUBLESHOOTING ============
doc.add_heading("11. Troubleshooting", level=1)
make_table(
    ["Symptom", "Cause", "Fix"],
    [["vllm: error: unrecognized arguments: --disable-log-requests", "Flag removed in vLLM 0.24",
      "Drop the flag (request logging is off by default)"],
     ["Warm TTFT identical to baseline in an LMCache mode", "Backend init failed; LMCache degrades to recompute and still serves correct answers",
      "grep 'init failed' in the server log; fix the YAML"],
     ["Engine dies during startup, no traceback, exit 137", "nixl_pool_size given in bytes → giant index → cgroup OOM kill",
      "Set nixl_pool_size to a slot count (e.g. 1500)"],
     ["OSError: Too many open files during pool creation", "One fd per pool slot vs default ulimit of 1024",
      "ulimit -n 65536 in the launch shell"],
     ["AssertionError: Invalid NIXL backend & device combination", "POSIX backend with nixl_buffer_device: cuda",
      "Use nixl_buffer_device: cpu for POSIX/HF3FS"],
     ["pip install fails: externally-managed-environment", "Image's system Python is PEP-668 locked",
      "uv pip install --python /opt/venv/bin/python3 <pkg>"],
     ["pkill kills your own SSH session", "pkill -f pattern matches the shell's own command line",
      "Use a bracket class: pkill -f 'vllm [s]erve'"],
     ["Disk fills up between rounds", "Disk-tier KV files (~53 GB) linger",
      "rm -rf /root/lmcache_disk/* before the NIXL round"],
     ["Next server launch fails: engine init error, GPU shows ~45 GB used with no server running",
      "vLLM engine cores rename themselves 'VLLM::EngineCore'; pkill 'vllm serve' orphans them holding GPU memory",
      "pkill -9 -f 'VLLM::' before relaunching (switch.sh does this)"],
     ["Server hangs at startup forever (futex wait) in a NIXL config", "use_direct_io: true deadlocks LMCache 0.5.1's scheduler-role init",
      "Remove use_direct_io from extra_config"],
     ["Fresh vast.ai instance shows 'exited' right after creation", "Host parked the container after the image pull",
      "vastai start instance <ID>, then wait for running + SSH"],
     ["Controller /lookup or token-targeted /clear always returns empty/0", "Chunk-key hashes differ across processes (randomized Python hash), or tokens lack the chat template",
      "export PYTHONHASHSEED=0 in engine AND controller; tokenize with messages + add_generation_prompt"]],
    [2.1, 2.4, 2.0], highlight_alt=True)

# ============ 12. INTERACTIVE DEMO ============
doc.add_heading("12. The interactive demo (dashboard)", level=1)
body("The repo also ships a browser dashboard for live, audience-facing demonstrations — the same stack "
     "as the benchmark, but driven by buttons with live gauges and on-page evidence. This section "
     "documents the current implementation.")

doc.add_heading("12.1 Architecture", level=2)
body("The control server (onbox_server.py) runs ON the GPU node, port 7811, and talks to vLLM over "
     "localhost; your browser reaches it through an auto-restarting SSH tunnel, so the page is only "
     "accessible from your machine. Every measurement is made on the box: ask.py sends each request to "
     "vLLM locally and reports server-side TTFT (your WAN latency is excluded by construction). Mode "
     "switches run detached via switch.sh (they survive SSH drops and kill orphaned engine cores); the "
     "LMCache controller (port 9050) provides pure per-tier eviction and the /lookup audit.", after=8)

doc.add_heading("12.2 Two serving modes", level=2)
make_table(
    ["Mode", "Topology", "What it teaches"],
    [["Baseline vLLM", "GPU 47 GB only", "Eviction = amnesia: any doc pushed out of HBM pays full prefill (~1.8 s) forever"],
     ["Full hierarchy", "GPU 47 GB → pinned CPU RAM 60 GB → NIXL pool 85.6 GB (NVMe)",
      "Write-through pyramid: each tier LRU-evicts at its cap while lower tiers retain copies; prefill is avoided while ANY tier holds the KV"]],
    [1.3, 2.6, 2.6])
body("Only two modes on purpose: the benchmark's single-tier modes isolate variables, but the demo "
     "isolates tiers with its clear buttons instead — clearing GPU shows the CPU tier serving, clearing "
     "GPU+CPU shows NIXL serving, clearing all three restores baseline behavior. All caps are printed on "
     "the page (GPU 47 GB is hardware truth; CPU 60 GB and the 1,700-slot x 48 MiB = 85.6 GB pool are "
     "deliberate config).", after=8)

doc.add_heading("12.3 The controls", level=2)
make_table(
    ["Control", "What it does"],
    [["Docs A-E · Cold prefill", "Single requests: five fixed ~10k-token docs (~2 GB KV each; docs 1-5 of set 1) and an always-new doc that can never be cached"],
     ["Sweep docs 1-10 / 11-20 / 21-30 / 31-40", "One manual pass per click over a 10-doc set (~20 GB KV, ~30 s). You compose passes yourself; 40 docs ≈ 80 GB exceeds every tier, forcing real eviction"],
     ["Clear GPU HBM", "Instant: vLLM /reset_prefix_cache (requires VLLM_SERVER_DEV_MODE=1)"],
     ["Clear CPU RAM tier", "Instant: controller /clear on LocalCPUBackend — returns a dropped-token receipt"],
     ["Clear NIXL/NVMe tier", "Instant: controller /clear on NixlStorageBackend (needs the clear() patch from 8.7)"],
     ["Clear ALL tiers", "All three at once with an itemized receipt — total amnesia in one click"],
     ["Your own prompt", "Free-form chat box (256-token replies); send a long paste twice to see caching on your own text"]],
    [1.9, 4.6])

doc.add_heading("12.4 Reading the numbers", level=2)
bullet("TTFT decoder: ~1.8 s cold prefill · ~0.07 s GPU hit · ~0.17 s CPU RAM · ~0.35 s NIXL/NVMe.",
       bold_lead="Which tier answered? ")
bullet("single requests end the clock at that request's server-side TTFT (wall time incl. network shown "
       "separately); batch sweeps keep the TOTAL batch time on the clock while per-request TTFTs land as "
       "individual rows in the persistent results log — TTFT is inherently a per-request metric.",
       bold_lead="Clock semantics: ")
bullet("host-resources row shows raw hardware (GPU used, RAM split into processes vs OS file cache, disk); "
       "the tier row shows KV residency per tier with caps — a model maintained from every request the "
       "dashboard serves, LRU-evicted at the true caps, persisted across restarts.",
       bold_lead="Gauges: ")

doc.add_heading("12.5 Click-for-evidence", level=2)
body("Clicking any tier card opens a terminal panel that re-runs real commands on the box at that moment "
     "(read-only, timestamped): nvidia-smi and vLLM's cache metrics for the GPU tier; for CPU and NIXL, "
     "the evidence_lookup.py audit — it tokenizes each demo doc exactly as the chat endpoint does "
     "(chat template + generation prompt) and asks the controller /lookup which tier holds those tokens. "
     "The audit totals match the gauges to within ~1%. The panel also states the known blind spot from "
     "8.7 (pool-only docs are invisible to /lookup in LMCache 0.5.1) and how to disprove it live.", after=8)

doc.add_heading("12.6 A six-step walkthrough for an audience", level=2)
body("① Sweep docs 1-10 — cold, ~1.8 s each; every gauge fills by 20 GB.  "
     "② Sweep 1-10 again — all ~0.07 s GPU hits.  "
     "③ Sweep 11-20, 21-30, 31-40 — 80 GB floods the pyramid; GPU pins at 100%, then CPU.  "
     "④ Sweep 1-10 again — the GPU evicted them long ago, yet no prefill: a lower tier answers.  "
     "⑤ Clear the CPU tier, sweep 11-20 — served from NIXL/NVMe at ~0.35 s, still no prefill.  "
     "⑥ Clear ALL tiers — the next sweep recomputes everything. That is baseline life, permanently.", after=8)

doc.add_heading("12.7 One-command deployment", level=2)
body("The deploy kit (~/kvcache/deploy/) automates this entire guide: rents a suitable A100, pushes all "
     "scripts, installs NIXL, downloads the model, applies the LMCache clear() patch, starts the "
     "controller and dashboard, launches full-hierarchy mode, and opens the tunnel:", after=4)
code(["cd ~/kvcache",
      "./deploy/deploy.sh                # rent cheapest suitable A100 and deploy (~5-8 min)",
      "./deploy/deploy.sh <INSTANCE_ID>  # or (re)deploy onto an existing box — idempotent",
      "# ends with: DEMO READY — open http://127.0.0.1:7811"])
body("See deploy/README.md for the full list of gotchas the kit bakes in (exited-on-create auto-start, "
     "hash-seed pinning, fd limits, the clear() patch, detached launches).", after=8)

# ============ 13. TEARDOWN ============
doc.add_heading("13. Tear down (do not skip)", level=1)
code(["vastai destroy instance <INSTANCE_ID>   # billing stops immediately",
      "vastai show instances                    # confirm the list is empty"])
callout("NOTE", "A stopped instance still bills for storage. Destroy, don't stop, when you are done. "
        "Everything needed to recreate the demo is in this guide and the scripts.")

# ============ APPENDIX ============
doc.add_page_break()
doc.add_heading("Appendix A — Full scripts", level=1)
body("Save each listing below as a plain-text file with the exact filename shown. A.1–A.4 go to /root/ "
     "on the instance (section 7.1); A.5 stays on your own machine (section 10).", italic=True)

for title, path in [("A.1  gen_docs.py — synthetic corpus generator", "/Users/wongtran/kvcache/remote/gen_docs.py"),
                    ("A.2  benchmark.py — TTFT benchmark client", "/Users/wongtran/kvcache/remote/benchmark.py"),
                    ("A.3  run_server.sh — four-mode server launcher", "/Users/wongtran/kvcache/remote/run_server.sh"),
                    ("A.4  run_mode.sh — one-command benchmark cycle", "/Users/wongtran/kvcache/remote/run_mode.sh"),
                    ("A.5  analyze.py — results comparison (runs on your machine)", "/Users/wongtran/kvcache/analyze.py")]:
    doc.add_heading(title, level=2)
    with open(path) as f:
        code(f.read().rstrip().split("\n"))

doc.save("/Users/wongtran/kvcache/report/KVCache-Tiering-Setup-Guide.docx")
print("saved KVCache-Tiering-Setup-Guide.docx")
